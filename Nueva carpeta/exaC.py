from docx import Document
from docx.shared import Cm, Pt


article_1 = """Bayern Munich came out on top in a thrilling German Cup final, beating Bayer Leverkusen 4-2 to secure its 20th title and remain on course for an historic treble.
David Alaba's stunning free kick and Serge Gnabry's clinical finish gave Bayern a commanding lead heading into half time and Hans-Dieter Flick's side seemingly already had one hand on the trophy.
However, Leverkusen responded well early in the second half and had a golden opportunity to halve the deficit through substitute Kevin Volland."""

article_2 = """(CNN)Liverpool got its Premier League title-winning celebrations back on track with a 2-0 win over Aston Villa, just days after being on the receiving end of a record-equaling defeat.
Many had suggested Jurgen Klopp's side was suffering from something of a hangover during Thursday's 4-0 demolition at the hands of Manchester City -- the joint-heaviest defeat by a team already crowned Premier League champion -- but Liverpool recovered in time to put relegation-threatened Aston Villa to the sword.
It wasn't all plain sailing at Anfield on Sunday as Villa wasted several good opportunities to take the lead, before Sadio Mane eventually broke the deadlock after 71 minutes. Villa, who gave the host a guard of honor before the game, had further chances to level the scores, but Liverpool youngster Curtis Jones wrapped up the victory in the dying moments with his first Premier League goal."""

list_of_articles = [article_1, article_2]

word_document = Document()
document_name = 'news-article-stats'

word_document = Document()
document_name = 'news-article-stats'

for article in list_of_articles:
    # extracting text stats
    text_stats = describe_text(article)
    text_stats['Article'] = article
    text_stats = dict(sorted(text_stats.items()))
    
    # customizing the table
    table = word_document.add_table(0, 0) # we add rows iteratively
    table.style = 'TableGrid'
    first_column_width = 5
    second_column_with = 10
    table.add_column(Cm(first_column_width))
    table.add_column(Cm(second_column_with))
    
    for index, stat_item in enumerate(text_stats.items()):
        table.add_row()
        stat_name, stat_result = stat_item
        row = table.rows[index]
        row.cells[0].text = str(stat_name)
        row.cells[1].text = str(stat_result)
    word_document.add_page_break()

word_document.save(document_name + '.docx')
